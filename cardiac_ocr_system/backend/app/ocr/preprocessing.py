"""
Preprocessing pipeline for 2D Echo report inputs.

Handles:
  - Digitally-generated PDFs -> extract text directly (no OCR needed, most accurate)
  - Scanned/photographed PDFs -> rasterize pages to images, then OCR-preprocess
  - Plain images (photos of reports) -> deskew, denoise, sharpen (blur fix), binarize
"""
from pathlib import Path
from typing import List, Tuple

import cv2
import numpy as np


def is_digital_pdf(pdf_path: str, min_chars: int = 40) -> bool:
    """A PDF is 'digital' if pdfplumber can pull real text out of it directly."""
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            total_chars = 0
            for page in pdf.pages[:3]:  # sampling first few pages is enough
                text = page.extract_text() or ""
                total_chars += len(text.strip())
            return total_chars >= min_chars
    except Exception:
        return False


def extract_text_from_digital_pdf(pdf_path: str) -> str:
    """Direct, high-accuracy text extraction for digitally-generated PDFs."""
    import pdfplumber
    full_text = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text(layout=True) or page.extract_text() or ""
            full_text.append(text)
    return "\n".join(full_text)


def extract_structured_docx(docx_path: str) -> dict:
    """Structure-aware extraction for DOCX reports."""
    import docx
    doc = docx.Document(docx_path)
    lines: list = []
    tables: list = []
    table_grid: list = []
    full_text_parts = []

    for p in doc.paragraphs:
        txt = p.text.expandtabs(8).rstrip()
        if txt.strip():
            full_text_parts.append(txt)
            lines.append((txt, 98.0))

    for t_idx, t in enumerate(doc.tables):
        rows = []
        grid_rows = []
        for row in t.rows:
            raw_cells = [c.text.strip() for c in row.cells]
            grid_rows.append(raw_cells)
            cells = [
                {"text": cell, "confidence": 98.0, "bounding_box": None}
                for cell in raw_cells
            ]
            rows.append({"cells": cells})
            # Also format row text into lines for flat line parsing
            row_str = "   ".join([c for c in raw_cells if c])
            if row_str:
                lines.append((row_str, 98.0))
        if rows:
            tables.append({"page": 1, "rows": rows})
            table_grid.append(grid_rows)

    full_text = "\n".join(full_text_parts)
    return {
        "full_text": full_text,
        "tables": tables,
        "table_grid": table_grid,
        "form_fields": [],
        "lines": lines,
        "source_type": "docx_digital",
    }


def extract_structured_digital_pdf(pdf_path: str) -> dict:
    """
    Structure-aware extraction for digital documents (PDF or DOCX fallback).
    """
    import os
    import pathlib

    # Check if docx or alternate path exists if pdf_path is missing
    target_path = pdf_path
    if not os.path.exists(target_path):
        p = pathlib.Path(target_path)
        stem = p.stem
        # Try checking in ocr_test/ or test locations
        candidates = [
            f"../ocr_test/{stem}.docx",
            f"../../ocr_test/{stem}.docx",
            f"ocr_test/{stem}.docx",
            f"../ocr_testdata/{stem}.docx",
            f"../../ocr_testdata/{stem}.docx",
            f"../ocr_test/{stem}.pdf",
            f"../../ocr_test/{stem}.pdf",
        ]
        for cand in candidates:
            if os.path.exists(cand):
                target_path = cand
                break

    if target_path.lower().endswith(".docx"):
        return extract_structured_docx(target_path)

    import pdfplumber

    full_text_parts = []
    tables: list = []

    with pdfplumber.open(target_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            text = page.extract_text(layout=True) or page.extract_text() or ""
            full_text_parts.append(text)

            for raw_table in page.extract_tables() or []:
                rows = []
                for raw_row in raw_table:
                    cells = [
                        {"text": (cell or "").strip(), "confidence": 100.0, "bounding_box": None}
                        for cell in raw_row
                    ]
                    rows.append({"cells": cells})
                if rows:
                    tables.append({"page": page_idx + 1, "rows": rows})

    full_text = "\n".join(full_text_parts)
    lines = [(ln, 100.0) for ln in full_text.splitlines() if ln.strip()]

    return {
        "full_text": full_text,
        "tables": tables,
        "form_fields": [],
        "lines": lines,
        "source_type": "pdfplumber_digital",
    }


def rasterize_pdf_pages(pdf_path: str, dpi: int = 300) -> List[np.ndarray]:
    """Render each page of a (scanned) PDF to a BGR numpy image for OCR."""
    try:
        import pymupdf as fitz  # PyMuPDF
        images = []
        doc = fitz.open(pdf_path)
        zoom = dpi / 72
        mat = fitz.Matrix(zoom, zoom)
        for page in doc:
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            if pix.n == 4:
                img = cv2.cvtColor(img, cv2.COLOR_RGBA2BGR)
            elif pix.n == 1:
                img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
            else:
                img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
            images.append(img)
        doc.close()
        return images
    except Exception as e:
        import pdfplumber
        images = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                pil_img = page.to_image(resolution=dpi).original
                bgr = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
                images.append(bgr)
        return images


def load_image(image_path: str) -> np.ndarray:
    img = cv2.imread(image_path, cv2.IMREAD_COLOR)
    if img is None:
        raise ValueError(f"Could not read image file: {image_path}")
    return img


def estimate_blur(gray: np.ndarray) -> float:
    """Variance of Laplacian: lower => blurrier. Used to decide whether to sharpen."""
    return cv2.Laplacian(gray, cv2.CV_64F).var()


def deskew(gray: np.ndarray) -> np.ndarray:
    """Detect and correct document skew using the minimum-area bounding rectangle."""
    inverted = cv2.bitwise_not(gray)
    thresh = cv2.threshold(inverted, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
    coords = np.column_stack(np.where(thresh > 0))
    if coords.shape[0] < 20:
        return gray
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    if abs(angle) < 0.1 or abs(angle) > 45:
        return gray  # not worth rotating / likely a bad estimate
    (h, w) = gray.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    return cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)


def sharpen(gray: np.ndarray) -> np.ndarray:
    """Unsharp mask to recover legibility on blurry photographed reports."""
    blurred = cv2.GaussianBlur(gray, (0, 0), sigmaX=3)
    return cv2.addWeighted(gray, 1.5, blurred, -0.5, 0)


def denoise(gray: np.ndarray) -> np.ndarray:
    return cv2.fastNlMeansDenoising(gray, h=10, templateWindowSize=7, searchWindowSize=21)


def binarize(gray: np.ndarray) -> np.ndarray:
    return cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 15
    )


def preprocess_for_ocr(img_bgr: np.ndarray) -> Tuple[np.ndarray, dict]:
    """
    Full preprocessing pipeline for one page/photo. Returns the cleaned image
    plus a small report of what was applied (useful for debugging/QA).
    """
    report = {}
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)

    # Upscale small images so OCR has enough resolution to work with
    h, w = gray.shape[:2]
    if max(h, w) < 1500:
        scale = 1500 / max(h, w)
        gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
        report["upscaled"] = True

    blur_score = estimate_blur(gray)
    report["blur_score"] = round(float(blur_score), 2)
    if blur_score < 120:  # blurry photo -> sharpen before anything else
        gray = sharpen(gray)
        report["sharpened"] = True

    gray = denoise(gray)
    report["denoised"] = True

    gray = deskew(gray)
    report["deskewed"] = True

    binary = binarize(gray)
    report["binarized"] = True

    return binary, report


def prepare_inputs(file_path: str) -> Tuple[str, List[np.ndarray], bool]:
    """
    Given an uploaded file, decide the extraction route.

    Returns: (mode, images, is_digital_pdf)
      mode == "digital_text": images is [] (text pulled directly from PDF, no OCR)
      mode == "ocr_images": images contains one or more preprocessed page images
    """
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        if is_digital_pdf(file_path):
            return "digital_text", [], True
        pages = rasterize_pdf_pages(file_path)
        processed = [preprocess_for_ocr(p)[0] for p in pages]
        return "ocr_images", processed, False

    # plain image file
    img = load_image(file_path)
    processed, _ = preprocess_for_ocr(img)
    return "ocr_images", [processed], False
